"""
ReportAgent — genera un reporte IR formal en DOCX después de una investigación.

Arquitectura (NIST 800-61 Rev 3 + DFIR-IRIS style):
  1. SQL determinista   → scope, timeline, IOCs, MITRE hits (sin LLM)
  2. qwen2.5:7b-instruct → executive summary + recommendations (2 secciones)
  3. python-docx        → ensambla el .docx final
  4. Auto-copia         → Kalishares si está montado

Secciones del reporte:
  0. Portada
  1. Resumen Ejecutivo   (LLM)
  2. Alcance del Incidente (SQL)
  3. Línea de Tiempo      (SQL — si hay timestamps)
  4. Análisis MITRE ATT&CK (MITRE rules + SQL)
  5. Indicadores de Compromiso (SQL)
  6. Hallazgos Clave      (EIL conclusion si existe)
  7. Recomendaciones      (LLM — 3 niveles)
  8. Gaps Forenses        (SQL + heurísticas)
"""

import json
import re
import sqlite3
import time
from datetime import datetime
from pathlib import Path

import httpx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI

REPORT_MODEL = "qwen2.5:7b-instruct"
OLLAMA_URL   = "http://localhost:11434"
TIMEOUT      = httpx.Timeout(connect=10.0, read=300.0, write=10.0, pool=5.0)
KALISHARES   = Path("/media/sf_Kalishares")

BOLD    = "\033[1m"
CYAN    = "\033[96m"
GREEN   = "\033[92m"
YELLOW  = "\033[93m"
DIM     = "\033[2m"
RESET   = "\033[0m"

# Colores MITRE para Word
_SEV_RGB = {
    "CRITICAL": RGBColor(0xC0, 0x00, 0x00),
    "HIGH":     RGBColor(0xFF, 0x66, 0x00),
    "MEDIUM":   RGBColor(0xFF, 0xC0, 0x00),
    "LOW":      RGBColor(0x00, 0x70, 0xC0),
}

# Event ID → descripción forense
_EID_NAMES = {
    4624: "Logon exitoso", 4625: "Logon fallido", 4634: "Logoff",
    4648: "Logon con credenciales explícitas", 4672: "Privilegios especiales asignados",
    4688: "Proceso creado", 4698: "Tarea programada creada",
    4720: "Cuenta de usuario creada", 4726: "Cuenta de usuario eliminada",
    4776: "Autenticación NTLM", 1102: "Log de seguridad borrado",
    4719: "Política de auditoría cambiada", 5140: "Recurso compartido accedido",
    5145: "Recurso compartido verificado", 4662: "Operación en objeto AD",
    5136: "Objeto AD modificado", 4732: "Miembro añadido a grupo local",
    4742: "Cuenta de equipo modificada", 1116: "Defender — malware detectado",
    1117: "Defender — acción tomada", 21: "RDP — logon exitoso",
    22: "RDP — shell iniciado", 23: "RDP — logoff", 24: "RDP — desconexión",
    1: "Sysmon — proceso creado", 3: "Sysmon — conexión de red",
    7: "Sysmon — imagen cargada", 10: "Sysmon — proceso accedido",
    11: "Sysmon — archivo creado", 13: "Sysmon — registro modificado",
    4771: "Kerberos — pre-auth fallida", 18456: "SQL Server — login fallido",
    4104: "PowerShell — script block logging",
}


# ── Fase 1: recolección determinista ─────────────────────────────────────────

def _collect_data(conn: sqlite3.Connection, case_name: str) -> dict:
    from nexus.stats import collect_basic_stats
    d = {"case_name": case_name, **collect_basic_stats(conn)}
    d["reg_keys"] = d.pop("registry_keys", 0)  # report usa reg_keys como alias

    def q(sql, params=()):
        try:
            return conn.execute(sql, params).fetchall()
        except Exception:
            return []

    def q1(sql, params=()):
        rows = q(sql, params)
        return rows[0][0] if rows else 0

    # Top event IDs con nombres
    rows = q("SELECT event_id, COUNT(*) n FROM events GROUP BY event_id ORDER BY n DESC LIMIT 15")
    d["top_event_ids"] = [
        {"event_id": r[0], "count": r[1], "name": _EID_NAMES.get(r[0], f"EID {r[0]}")}
        for r in rows
    ]

    # Usuarios humanos (sin cuentas de sistema)
    rows = q(
        "SELECT username, COUNT(*) n FROM events "
        "WHERE username IS NOT NULL AND username != '' "
        "AND username NOT IN ('SYSTEM','LOCAL SERVICE','NETWORK SERVICE','ANONYMOUS LOGON') "
        "AND username NOT LIKE '%$' AND username NOT LIKE 'NT AUTHORITY%' "
        "GROUP BY username ORDER BY n DESC LIMIT 10"
    )
    d["human_users"] = [{"username": r[0], "count": r[1]} for r in rows]

    # Máquinas
    rows = q("SELECT computer, COUNT(*) n FROM events WHERE computer IS NOT NULL AND computer != '' GROUP BY computer ORDER BY n DESC LIMIT 10")
    d["machines"] = [{"computer": r[0], "count": r[1]} for r in rows]

    # Timeline de eventos clave (con timestamps)
    if d["has_timestamps"]:
        rows = q(
            "SELECT timestamp_utc, event_id, username, source_ip, computer "
            "FROM events WHERE timestamp_utc IS NOT NULL AND timestamp_utc != '' "
            "ORDER BY timestamp_utc ASC LIMIT 20"
        )
        d["timeline_events"] = [
            {"ts": r[0][:19], "eid": r[1], "user": r[2] or "", "ip": r[3] or "", "computer": r[4] or ""}
            for r in rows
        ]
    else:
        d["timeline_events"] = []

    # IOC — IPs externas con actividad
    rows = q(
        "SELECT source_ip, COUNT(*) n, "
        "SUM(CASE WHEN event_id IN (4624,21) THEN 1 ELSE 0 END) logons_ok, "
        "SUM(CASE WHEN event_id IN (4625,4771,4776,18456) THEN 1 ELSE 0 END) logons_fail "
        "FROM events "
        "WHERE source_ip IS NOT NULL AND source_ip != '' "
        "AND source_ip NOT LIKE '10.%' AND source_ip NOT LIKE '192.168.%' "
        "AND source_ip NOT LIKE '127.%' AND source_ip NOT LIKE '172.1%' "
        "GROUP BY source_ip ORDER BY n DESC LIMIT 10"
    )
    d["external_ips"] = [
        {"ip": r[0], "total": r[1], "logons_ok": r[2], "logons_fail": r[3]}
        for r in rows
    ]

    # IOC — Procesos sospechosos
    rows = q(
        "SELECT name, exe_path, username, command_line FROM processes "
        "WHERE exe_path LIKE '%\\Temp\\%' OR exe_path LIKE '%\\AppData\\%' "
        "OR exe_path LIKE '%\\Users\\Public\\%' "
        "OR LOWER(name) IN ('mimikatz.exe','procdump.exe','wce.exe','mshta.exe') "
        "LIMIT 10"
    )
    d["suspicious_procs"] = [
        {"name": r[0], "path": (r[1] or "")[:60], "user": r[2] or "", "cmd": (r[3] or "")[:60]}
        for r in rows
    ]

    # IOC — Conexiones externas establecidas
    rows = q(
        "SELECT remote_address, remote_port, protocol, state, process_name "
        "FROM network_connections "
        "WHERE state='ESTABLISHED' AND remote_address IS NOT NULL AND remote_address != '' "
        "AND remote_address NOT LIKE '10.%' AND remote_address NOT LIKE '192.168.%' "
        "AND remote_address NOT LIKE '127.%' LIMIT 10"
    )
    d["external_conns"] = [
        {"ip": r[0], "port": r[1], "proto": r[2], "state": r[3], "proc": r[4] or ""}
        for r in rows
    ]

    # IOC — Registry Run keys
    rows = q(
        "SELECT key_path, value_name, value_data FROM registry_keys "
        "WHERE key_path LIKE '%\\Run%' OR key_path LIKE '%\\RunOnce%' LIMIT 10"
    )
    d["run_keys"] = [
        {"path": (r[0] or "")[:70], "name": r[1] or "", "data": (r[2] or "")[:60]}
        for r in rows
    ]

    # Tareas sospechosas
    rows = q(
        "SELECT task_name, command, author, run_as FROM scheduled_tasks "
        "WHERE command LIKE '%Temp%' OR command LIKE '%AppData%' "
        "OR command LIKE '%-Enc%' OR command LIKE '%DownloadString%' LIMIT 5"
    )
    d["suspicious_tasks"] = [
        {"name": r[0] or "", "cmd": (r[1] or "")[:60], "author": r[2] or "", "run_as": r[3] or ""}
        for r in rows
    ]

    # MITRE ATT&CK
    try:
        from nexus.router import tool_threat_hunt
        from nexus.finding_validator import enrich_hits
        hits = tool_threat_hunt(conn)
        if hits:
            hits = enrich_hits(hits)
        d["mitre_hits"] = [
            {
                "rule_id": h["rule_id"],
                "severity": h["severity"],
                "name": h["name"],
                "count": h["count"],
                "table": h["table"],
                "confidence": h["validation"].confidence if "validation" in h else 0.5,
                "risk_label": h["validation"].risk_label if "validation" in h else "POSSIBLE",
                "fp_risk": h["validation"].fp_risk if "validation" in h else "medium",
            }
            for h in hits
        ]
    except Exception:
        d["mitre_hits"] = []

    # Evidence files list
    rows = q("SELECT filename, evidence_type, file_size_kb, record_count, ingested_at FROM evidence_files ORDER BY ingested_at LIMIT 20")
    d["evidence_list"] = [
        {"file": r[0], "type": r[1], "size_kb": round(r[2] or 0, 1), "records": r[3] or 0}
        for r in rows
    ]

    return d


# ── Fase 2: LLM — solo 2 secciones ───────────────────────────────────────────

def _build_llm_context(data: dict, eil_conclusion: str) -> str:
    lines = [f"Caso forense: {data['case_name']}"]
    lines.append(f"Eventos totales: {data['total_events']:,} | Usuarios: {data['unique_users']} | IPs: {data['unique_ips']} | Máquinas: {data['unique_machines']}")

    if data["has_timestamps"]:
        lines.append(f"Período: {data['first_event'][:10]} → {data['last_event'][:10]} ({data['dwell_days']} días)")

    if data["mitre_hits"]:
        lines.append(f"\nReglas MITRE disparadas ({len(data['mitre_hits'])}):")
        for h in data["mitre_hits"][:8]:
            lines.append(f"  [{h['severity']}] {h['rule_id']} — {h['name']}: {h['count']} hits ({h['risk_label']})")

    if data["external_ips"]:
        lines.append("\nIPs externas con actividad:")
        for ip in data["external_ips"][:5]:
            lines.append(f"  {ip['ip']}: {ip['total']} eventos, {ip['logons_ok']} logons exitosos, {ip['logons_fail']} fallidos")

    if data["brute_force_total"] > 0:
        lines.append(f"\nFuerza bruta: {data['brute_force_total']:,} eventos de autenticación fallida")

    if data["log_clearing"] > 0:
        lines.append(f"Borrado de logs: {data['log_clearing']} eventos (evasión de defensa)")

    if data["lsass_access"] > 0:
        lines.append(f"Acceso a LSASS: {data['lsass_access']} eventos (volcado de credenciales)")

    if eil_conclusion:
        lines.append(f"\nConclusión del agente de investigación (EIL):\n{eil_conclusion}")

    return "\n".join(lines)


def _llm_executive_summary(context: str, model: str) -> str:
    client = OpenAI(base_url=f"{OLLAMA_URL}/v1", api_key="ollama", timeout=TIMEOUT, max_retries=0)
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": (
                "Eres un analista DFIR senior redactando el resumen ejecutivo de un reporte de respuesta a incidentes. "
                "Escribe en español, tono profesional, sin tecnicismos innecesarios. "
                "Máximo 4 párrafos. Sin encabezados, sin bullet points. Solo texto corrido. "
                "Incluye: qué pasó, el alcance del impacto, la fase más avanzada del atacante, y el estado actual."
            )},
            {"role": "user", "content": f"Genera el resumen ejecutivo para este caso:\n\n{context}"},
        ],
        temperature=0.2,
        max_tokens=400,
    )
    return (resp.choices[0].message.content or "").strip()


def _llm_recommendations(context: str, model: str) -> str:
    client = OpenAI(base_url=f"{OLLAMA_URL}/v1", api_key="ollama", timeout=TIMEOUT, max_retries=0)
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": (
                "Eres un analista DFIR senior. Genera recomendaciones de remediación en español. "
                "Estructura EXACTA requerida — usa exactamente estos encabezados:\n"
                "INMEDIATO (0-24 horas):\n- acción 1\n- acción 2\n\n"
                "CORTO PLAZO (1-4 semanas):\n- acción 1\n- acción 2\n\n"
                "ESTRATÉGICO (1-6 meses):\n- acción 1\n- acción 2\n"
                "Máximo 3 acciones por nivel. Específicas y accionables."
            )},
            {"role": "user", "content": f"Genera recomendaciones para este caso:\n\n{context}"},
        ],
        temperature=0.2,
        max_tokens=400,
    )
    return (resp.choices[0].message.content or "").strip()


# ── Fase 3: ensamblado DOCX ───────────────────────────────────────────────────

def _add_heading(doc, text: str, level: int):
    doc.add_heading(text, level=level)


def _add_para(doc, text: str, bold: bool = False, size: int = 11, color: RGBColor = None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def _add_table(doc, headers: list, rows: list, col_widths: list = None):
    if not rows:
        doc.add_paragraph("Sin datos.", style="Normal")
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    # Data
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val) if val is not None else ""
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Inches(w)


def _build_docx(data: dict, exec_summary: str, recommendations: str, eil_conclusion: str, triage_result: dict) -> Document:
    doc = Document()

    # ── Portada ───────────────────────────────────────────────────────────────
    doc.add_heading("REPORTE DE RESPUESTA A INCIDENTE", 0)
    doc.add_heading(f"Caso: {data['case_name']}", 1)

    meta = doc.add_paragraph()
    meta.add_run(f"Fecha del reporte: {datetime.now().strftime('%Y-%m-%d %H:%M')} UTC\n")
    meta.add_run(f"Generado por: Nexus DFIR (air-gap)\n")
    if data["has_timestamps"]:
        meta.add_run(f"Período del incidente: {data['first_event'][:10]} → {data['last_event'][:10]}\n")
    meta.add_run("TLP: AMBER — Distribución restringida\n")
    meta.add_run("Clasificación: CONFIDENCIAL")

    if triage_result:
        sev = triage_result.get("severity", "DESCONOCIDA")
        meta.add_run(f"\nSeveridad del triage: {sev}")

    doc.add_page_break()

    # ── 1. Resumen Ejecutivo ──────────────────────────────────────────────────
    _add_heading(doc, "1. Resumen Ejecutivo", 1)
    if exec_summary:
        doc.add_paragraph(exec_summary)
    else:
        doc.add_paragraph("Resumen no disponible.")
    doc.add_paragraph()

    # ── 2. Alcance del Incidente ──────────────────────────────────────────────
    _add_heading(doc, "2. Alcance del Incidente", 1)

    _add_heading(doc, "2.1 Estadísticas Generales", 2)
    _add_table(doc,
        ["Métrica", "Valor"],
        [
            ["Total de eventos",        f"{data['total_events']:,}"],
            ["Event IDs únicos",         str(data["unique_event_ids"] if "unique_event_ids" in data else "—")],
            ["Usuarios únicos",          str(data["unique_users"])],
            ["IPs de origen únicas",     str(data["unique_ips"])],
            ["Máquinas involucradas",    str(data["unique_machines"])],
            ["Procesos capturados",      str(data["processes"])],
            ["Conexiones de red",        str(data["net_connections"])],
            ["Tareas programadas",       str(data["sched_tasks"])],
            ["Claves de registro",       str(data["reg_keys"])],
            ["Archivos de evidencia",    str(data["evidence_files"])],
        ] + ([["Período del incidente", f"{data['first_event'][:10]} → {data['last_event'][:10]} ({data['dwell_days']} días)"]] if data["has_timestamps"] else []),
        col_widths=[2.5, 3.5]
    )
    doc.add_paragraph()

    if data["human_users"]:
        _add_heading(doc, "2.2 Cuentas de Usuario", 2)
        _add_table(doc,
            ["Usuario", "Eventos"],
            [[u["username"], f"{u['count']:,}"] for u in data["human_users"]],
            col_widths=[3.5, 2.5]
        )
        doc.add_paragraph()

    if data["machines"]:
        _add_heading(doc, "2.3 Sistemas Involucrados", 2)
        _add_table(doc,
            ["Hostname", "Eventos"],
            [[m["computer"], f"{m['count']:,}"] for m in data["machines"]],
            col_widths=[3.5, 2.5]
        )
        doc.add_paragraph()

    if data["top_event_ids"]:
        _add_heading(doc, "2.4 Distribución de Eventos", 2)
        _add_table(doc,
            ["Event ID", "Nombre", "Cantidad"],
            [[str(e["event_id"]), e["name"], f"{e['count']:,}"] for e in data["top_event_ids"]],
            col_widths=[1.0, 4.0, 1.5]
        )
        doc.add_paragraph()

    # ── 3. Línea de Tiempo ────────────────────────────────────────────────────
    if data["has_timestamps"] and data["timeline_events"]:
        _add_heading(doc, "3. Línea de Tiempo del Incidente", 1)
        _add_table(doc,
            ["Timestamp (UTC)", "Event ID", "Usuario", "IP Origen", "Equipo"],
            [
                [e["ts"], str(e["eid"]), e["user"], e["ip"], e["computer"]]
                for e in data["timeline_events"]
            ],
            col_widths=[1.8, 0.8, 1.4, 1.4, 1.1]
        )
        doc.add_paragraph()
    else:
        _add_heading(doc, "3. Línea de Tiempo del Incidente", 1)
        doc.add_paragraph("La evidencia disponible no contiene timestamps procesables. La reconstrucción cronológica requiere correlación manual de artefactos adicionales (MFT, prefetch, registry).")
        doc.add_paragraph()

    # ── 4. Análisis MITRE ATT&CK ─────────────────────────────────────────────
    _add_heading(doc, "4. Análisis MITRE ATT&CK", 1)

    if data["mitre_hits"]:
        _add_table(doc,
            ["Técnica", "Severidad", "Nombre", "Hits", "Confianza", "FP Risk"],
            [
                [
                    h["rule_id"],
                    h["severity"],
                    h["name"],
                    str(h["count"]),
                    h["risk_label"],
                    h["fp_risk"],
                ]
                for h in sorted(data["mitre_hits"], key=lambda x: {"CRITICAL":0,"HIGH":1,"MEDIUM":2,"LOW":3}.get(x["severity"],4))
            ],
            col_widths=[1.0, 0.8, 2.5, 0.6, 1.0, 0.7]
        )
    else:
        doc.add_paragraph("No se activaron reglas MITRE ATT&CK en este caso. La ausencia de alertas no descarta la presencia de actividad maliciosa — puede indicar artefactos insuficientes o técnicas fuera del conjunto de reglas evaluado.")
    doc.add_paragraph()

    # ── 5. Indicadores de Compromiso (IOC) ────────────────────────────────────
    _add_heading(doc, "5. Indicadores de Compromiso (IOC)", 1)

    if data["external_ips"]:
        _add_heading(doc, "5.1 IPs Externas con Actividad", 2)
        _add_table(doc,
            ["IP", "Total Eventos", "Logons Exitosos", "Logons Fallidos"],
            [[ip["ip"], str(ip["total"]), str(ip["logons_ok"]), str(ip["logons_fail"])] for ip in data["external_ips"]],
            col_widths=[1.8, 1.2, 1.5, 1.5]
        )
        doc.add_paragraph()

    if data["suspicious_procs"]:
        _add_heading(doc, "5.2 Procesos Sospechosos", 2)
        _add_table(doc,
            ["Nombre", "Ruta", "Usuario", "Comando"],
            [[p["name"], p["path"], p["user"], p["cmd"]] for p in data["suspicious_procs"]],
            col_widths=[1.2, 2.0, 1.2, 2.1]
        )
        doc.add_paragraph()

    if data["external_conns"]:
        _add_heading(doc, "5.3 Conexiones Externas Activas", 2)
        _add_table(doc,
            ["IP Remota", "Puerto", "Protocolo", "Estado", "Proceso"],
            [[c["ip"], str(c["port"]), c["proto"] or "", c["state"] or "", c["proc"]] for c in data["external_conns"]],
            col_widths=[1.8, 0.7, 0.8, 0.9, 2.3]
        )
        doc.add_paragraph()

    if data["run_keys"]:
        _add_heading(doc, "5.4 Registry Run Keys (Persistencia)", 2)
        _add_table(doc,
            ["Clave", "Nombre", "Valor"],
            [[r["path"], r["name"], r["data"]] for r in data["run_keys"]],
            col_widths=[2.5, 1.5, 2.5]
        )
        doc.add_paragraph()

    if data["suspicious_tasks"]:
        _add_heading(doc, "5.5 Tareas Programadas Sospechosas", 2)
        _add_table(doc,
            ["Nombre", "Comando", "Autor", "Ejecuta como"],
            [[t["name"], t["cmd"], t["author"], t["run_as"]] for t in data["suspicious_tasks"]],
            col_widths=[1.5, 2.5, 1.2, 1.3]
        )
        doc.add_paragraph()

    if not any([data["external_ips"], data["suspicious_procs"], data["external_conns"], data["run_keys"], data["suspicious_tasks"]]):
        doc.add_paragraph("No se identificaron IOCs mediante las heurísticas automáticas. Revisar evidencia manualmente.")
        doc.add_paragraph()

    # ── 6. Hallazgos Clave (EIL) ──────────────────────────────────────────────
    _add_heading(doc, "6. Hallazgos Clave", 1)
    if eil_conclusion:
        doc.add_paragraph("Los siguientes hallazgos fueron determinados por el agente de investigación autónomo (EIL — Evidence Interrogation Loop):")
        doc.add_paragraph(eil_conclusion)
    else:
        doc.add_paragraph("No se ejecutó una investigación autónoma (EIL) para este caso.")
        doc.add_paragraph("Para obtener una narrativa completa del incidente, ejecutar:")
        doc.add_paragraph(f'  nexus investigate {data["case_name"]} "¿Qué pasó en este incidente?"')
    doc.add_paragraph()

    # ── 7. Recomendaciones ────────────────────────────────────────────────────
    _add_heading(doc, "7. Recomendaciones", 1)
    if recommendations:
        for line in recommendations.splitlines():
            line = line.strip()
            if not line:
                continue
            if line.endswith(":") and line.isupper() or "(" in line and line.endswith("):"):
                _add_heading(doc, line, 2)
            elif line.startswith("- ") or line.startswith("• "):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(line[2:])
            else:
                doc.add_paragraph(line)
    else:
        doc.add_paragraph("Recomendaciones no disponibles.")
    doc.add_paragraph()

    # ── 8. Gaps Forenses ──────────────────────────────────────────────────────
    _add_heading(doc, "8. Gaps Forenses", 1)
    gaps = []
    if not data["has_timestamps"]:
        gaps.append("Sin timestamps en eventos — reconstrucción cronológica limitada. Recuperar EVTX originales con herramientas como EvtxECmd.")
    if data["processes"] == 0:
        gaps.append("Sin snapshot de procesos — no es posible analizar la cadena de ejecución en vivo. Agregar tasklist/WMIC/Sysmon EID 1.")
    if data["net_connections"] == 0:
        gaps.append("Sin datos de red en tiempo real — actividad de C2 activa no confirmable. Agregar netstat/PCAP.")
    if data["sched_tasks"] == 0:
        gaps.append("Sin tareas programadas — persistencia vía schtasks no confirmable. Exportar con 'schtasks /fo CSV'.")
    if data["reg_keys"] == 0:
        gaps.append("Sin exportación de registro — persistencia vía Run keys no confirmable. Exportar HKLM\\SOFTWARE\\Microsoft\\Windows\\CurrentVersion\\Run.")
    if not eil_conclusion:
        gaps.append("No se ejecutó investigación autónoma (EIL) — la narrativa completa del ataque no ha sido construida automáticamente.")
    if data["lsass_access"] == 0 and data["brute_force_total"] > 10:
        gaps.append("Brute force detectado pero sin acceso confirmado a LSASS — volcado de credenciales no confirmado. Agregar Sysmon EID 10.")
    if not gaps:
        gaps.append("No se identificaron gaps significativos en la evidencia disponible para este caso.")

    for gap in gaps:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(gap)

    doc.add_paragraph()

    # ── 9. Evidencia Analizada ────────────────────────────────────────────────
    if data["evidence_list"]:
        _add_heading(doc, "9. Evidencia Analizada", 1)
        _add_table(doc,
            ["Archivo", "Tipo", "Tamaño (KB)", "Registros"],
            [[e["file"], e["type"], str(e["size_kb"]), str(e["records"])] for e in data["evidence_list"]],
            col_widths=[2.5, 1.5, 1.0, 1.0]
        )
        doc.add_paragraph()

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"Generado automáticamente por Nexus DFIR v0.2.0 — {datetime.now().strftime('%Y-%m-%d %H:%M UTC')}").italic = True
    p.add_run("\ngithub.com/robdinovil/nexus-dfir  |  100% air-gap  |  CPU-only").italic = True

    return doc


# ── Punto de entrada ──────────────────────────────────────────────────────────

def _build_markdown(data: dict, exec_summary: str, recommendations: str,
                    eil_conclusion: str, triage_result: dict) -> str:
    """Genera el reporte IR completo en Markdown."""
    now = datetime.now().strftime("%Y-%m-%d %H:%M UTC")
    lines = [
        f"# Reporte de Respuesta a Incidentes — {data['case_name']}",
        f"**Generado**: {now} | **Herramienta**: Nexus DFIR v0.2.0 (air-gap, CPU-only)",
        "",
        "---",
        "",
        "## 1. Resumen Ejecutivo",
        "",
        exec_summary,
        "",
        "---",
        "",
        "## 2. Alcance del Incidente",
        "",
        f"| Métrica | Valor |",
        f"|---|---|",
        f"| Eventos totales | {data['total_events']:,} |",
        f"| Usuarios únicos | {data['unique_users']} |",
        f"| IPs únicas | {data['unique_ips']} |",
        f"| Máquinas | {data['unique_machines']} |",
        f"| Procesos | {data['processes']:,} |",
        f"| Conexiones de red | {data['net_connections']:,} |",
        f"| Tareas programadas | {data['sched_tasks']} |",
        f"| Claves de registro | {data.get('reg_keys', 0)} |",
    ]
    if data.get("has_timestamps"):
        lines += [
            f"| Inicio del período | {data['first_event'][:19]} |",
            f"| Fin del período | {data['last_event'][:19]} |",
            f"| Dwell time | {data['dwell_days']} días |",
        ]
    lines += [""]

    if data.get("mitre_hits"):
        lines += ["---", "", "## 3. Análisis MITRE ATT&CK", "",
                  "| Regla | Severidad | Hits | Confianza | Riesgo FP |",
                  "|---|---|---|---|---|"]
        for h in data["mitre_hits"]:
            lines.append(f"| {h['rule_id']} — {h['name']} | {h['severity']} | {h['count']} | {h.get('confidence', 0):.2f} | {h.get('fp_risk', '?')} |")
        lines.append("")

    if data.get("external_ips"):
        lines += ["---", "", "## 4. Indicadores de Compromiso", "", "### IPs Externas", ""]
        for ip in data["external_ips"]:
            lines.append(f"- `{ip['ip']}` — {ip['total']} eventos, {ip['logons_ok']} logons OK, {ip['logons_fail']} fallidos")
        lines.append("")

    if eil_conclusion:
        lines += ["---", "", "## 5. Hallazgos Clave (EIL)", "", eil_conclusion, ""]

    lines += ["---", "", "## 6. Recomendaciones", "", recommendations, ""]

    triage_sev = triage_result.get("severity", "")
    if triage_sev:
        lines += ["---", "", "## 7. Clasificación de Triage", "",
                  f"**Severidad**: {triage_sev}",
                  f"**Fase**: {triage_result.get('attack_phase', '')}",
                  ""]

    lines += ["---", "", f"*Nexus DFIR — github.com/robdinovil/nexus-dfir*", ""]
    return "\n".join(lines)


def report(
    case_name: str,
    db_path: str,
    case_dir: str,
    eil_conclusion: str = "",
    triage_json: str = "",
    model: str = REPORT_MODEL,
    fmt: str = "docx",
    verbose: bool = True,
) -> str:
    """
    Genera el reporte IR para un caso Nexus.
    fmt: 'docx' (default) | 'md' — formato de salida.
    Devuelve la ruta al archivo generado.
    """
    t0 = time.time()

    def _print(msg):
        if verbose:
            print(msg, flush=True)

    _print(f"\n  {DIM}[REPORT] Recopilando datos del caso...{RESET}")

    conn = sqlite3.connect(db_path, check_same_thread=False)
    data = _collect_data(conn, case_name)
    # Inject unique_event_ids (collected separately since _collect_data may not have it)
    data["unique_event_ids"] = conn.execute("SELECT COUNT(DISTINCT event_id) FROM events").fetchone()[0]
    conn.close()

    # Cargar triage previo si existe
    triage_result = {}
    if triage_json and Path(triage_json).exists():
        try:
            triage_result = json.loads(Path(triage_json).read_text())
        except Exception:
            pass
    else:
        # Buscar el triage más reciente en case_dir
        triage_files = sorted(Path(case_dir).glob("triage_*.json"), reverse=True)
        if triage_files:
            try:
                triage_result = json.loads(triage_files[0].read_text())
                _print(f"  {DIM}[REPORT] Usando triage: {triage_files[0].name}{RESET}")
            except Exception:
                pass

    # Construir contexto para el LLM
    context = _build_llm_context(data, eil_conclusion)

    _print(f"  {DIM}[REPORT] Generando resumen ejecutivo ({model})...{RESET}")
    exec_summary = _llm_executive_summary(context, model)

    _print(f"  {DIM}[REPORT] Generando recomendaciones ({model})...{RESET}")
    recommendations = _llm_recommendations(context, model)

    ts_str = datetime.now().strftime("%Y%m%d_%H%M%S")

    if fmt == "md":
        _print(f"  {DIM}[REPORT] Ensamblando Markdown...{RESET}")
        md_content = _build_markdown(data, exec_summary, recommendations, eil_conclusion, triage_result)
        out_name = f"IR_{case_name}_{ts_str}.md"
        out_path = Path(case_dir) / out_name
        out_path.write_text(md_content, encoding="utf-8")
    else:
        _print(f"  {DIM}[REPORT] Ensamblando DOCX...{RESET}")
        doc = _build_docx(data, exec_summary, recommendations, eil_conclusion, triage_result)
        out_name = f"IR_{case_name}_{ts_str}.docx"
        out_path = Path(case_dir) / out_name
        doc.save(str(out_path))

    elapsed = time.time() - t0

    _print(f"\n  {BOLD}{GREEN}✓ Reporte generado en {elapsed:.0f}s{RESET}")
    _print(f"  {GREEN}  {out_path}{RESET}")

    # Copiar a Kalishares si está montado
    if KALISHARES.exists():
        import shutil
        ks_path = KALISHARES / out_name
        shutil.copy2(str(out_path), str(ks_path))
        _print(f"  {GREEN}  Copiado a Kalishares: {ks_path.name}{RESET}")

    _print("")
    return str(out_path)
